import streamlit as st
import os
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
import yfinance as yf
import requests
import time
from datetime import datetime, timedelta
from io import BytesIO
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import plotly.graph_objects as go
from rich.console import Console
from rich.table import Table

# --- [1. 환경 변수 및 OpenAI 설정] ---
load_dotenv()
api_key = os.getenv("Open_api_key")
client = OpenAI(api_key=api_key)

# --- [2. 페이지 기본 설정] ---
st.set_page_config(page_title="Trade Master 2026", layout="wide", page_icon="🚢")

# Matplotlib 및 전체 폰트 통일 설정 (Pretendard 기반)
plt.rcParams['font.family'] = 'Pretendard'
plt.rcParams['axes.unicode_minus'] = False

# --- [3. UI 디자인 및 스타일링 (Pretendard 폰트 통일 및 오류 수정)] ---
st.markdown("""
    <style>
    @import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');
    
    /* 전체 폰트 통일 */
    html, body, [class*="css"], .stApp, * {
        font-family: 'Pretendard', -apple-system, BlinkMacSystemFont, system-ui, Roboto, sans-serif !important;
    }
    
    .stApp { background-color: #ffffff; color: #31333f; }
    section[data-testid="stSidebar"] { background-color: #f0f2f6 !important; }
    
    /* 버튼 스타일 */
    .stButton>button { width: 100%; border-radius: 8px; height: 3em; background-color: #3d5afe; color: white; border: none; transition: 0.3s; font-weight: 600; }
    .stButton>button:hover { background-color: #1a237e; box-shadow: 0 4px 12px rgba(61, 90, 254, 0.3); }
    
    /* 안내 박스 */
    .info-box { background-color: #f8faff; padding: 20px; border-radius: 12px; border-left: 6px solid #3d5afe; color: #1a237e; font-weight: 500; box-shadow: 0 2px 10px rgba(0,0,0,0.05); }
    
    /* 환율 표 디자인 (Pretendard 적용) */
    .modern-table-container {
        margin: 20px 0;
        border-radius: 12px;
        overflow: hidden;
        border: 1px solid #e0e6ed;
        box-shadow: 0 4px 20px rgba(0, 0, 0, 0.05);
    }
    .modern-table {
        width: 100%;
        border-collapse: collapse;
        background-color: white;
    }
    .modern-table th {
        background-color: #f1f4f9;
        color: #475569;
        font-weight: 600;
        padding: 16px;
        text-align: center;
        border-bottom: 2px solid #e2e8f0;
        font-size: 14px;
    }
    .modern-table td {
        padding: 14px 20px;
        border-bottom: 1px solid #f1f5f9;
        font-size: 15px;
        color: #1e293b;
    }
    .modern-table tr:last-child td { border-bottom: none; }
    .modern-table tr:hover { background-color: #f8fafc; }
    
    .currency-name { font-weight: 700; color: #0f172a; text-align: left !important; }
    .rate-val { font-weight: 600; text-align: right !important; color: #2563eb; letter-spacing: -0.02em; }
    .rate-diff { text-align: right !important; font-size: 13px; font-weight: 500; }
    .diff-up { color: #dc2626; } 
    .diff-down { color: #2563eb; } 

    /* 사이드바 팝오버 버튼 스타일 및 expand_more 글자 제거 */
    div[data-testid="stSidebar"] .stPopover > button {
        background-color: #ff9800 !important;
        color: white !important;
        border: none;
        font-weight: 600;
    }
    /* 팝오버 내부의 화살표 아이콘 및 expand_more 텍스트 강제 숨기기 */
    div[data-testid="stSidebar"] .stPopover button svg,
    div[data-testid="stSidebar"] .stPopover button span:last-child {
        display: none !important;
    }
    </style>
    """, unsafe_allow_html=True)

# --- [4. 세션 상태 초기화] ---
if 'exchange_rates' not in st.session_state:
    st.session_state['exchange_rates'] = {"USD": 1440.70, "JPY": 935.94, "EUR": 1717.31, "CNY": 207.38}
if 'use_realtime' not in st.session_state:
    st.session_state['use_realtime'] = False

# --- [5. 환율 관련 함수] ---
def get_realtime_exchange_rates():
    """yfinance를 사용하여 실시간 환율 정보를 가져옵니다."""
    tickers = {"USD": "USDKRW=X", "JPY": "JPYKRW=X", "EUR": "EURKRW=X", "CNY": "CNYKRW=X"}
    updated_rates = {}
    try:
        for code, ticker in tickers.items():
            data = yf.download(ticker, period="2d", interval="1d", progress=False)
            if not data.empty:
                val = data['Close'].iloc[-1]
                updated_rates[code] = float(val) * 100 if code == "JPY" else float(val)
            else:
                updated_rates[code] = st.session_state['exchange_rates'][code]
        return updated_rates
    except Exception as e:
        st.error(f"데이터를 가져오는 중 오류 발생: {e}")
        return st.session_state['exchange_rates']

@st.cache_data(ttl=3600)
def get_currency_history(ticker_symbol, base_val, multiplier, use_realtime, current_date):
    if use_realtime:
        try:
            data = yf.download(ticker_symbol, period="1mo", interval="1d", progress=False)
            if not data.empty and not data['Close'].isnull().all():
                df = data[['Close']].reset_index()
                df.columns = ["날짜", "환율"]
                df['환율'] = df['환율'] * multiplier
                df['날짜'] = pd.to_datetime(df['날짜']).dt.date
                return df.sort_values(by="날짜")
        except Exception: pass
    
    np.random.seed(abs(hash(ticker_symbol)) % (10**8))
    dates = pd.date_range(end=current_date, periods=30)
    values = base_val + np.cumsum(np.random.randn(30) * (base_val * 0.005))
    return pd.DataFrame({"날짜": dates.date, "환율": values})

def calculate_estimated_cost(base_price, term, transport, insurance, payment, fta_type):
    total = base_price
    if term in ["CFR", "CIF", "CPT", "CIP", "DAP", "DPU", "DDP"]:
        freight_rate = 0.15 if transport == "항공(AIR)" else 0.05
        total += base_price * freight_rate
    ins_rates = {"ICC(A) (=ICC(AIR))": 0.008, "ICC(B)": 0.005, "ICC(C)": 0.003, "선택 안함": 0}
    total += base_price * ins_rates.get(insurance, 0)
    pay_fees = {"사전 송금": 0.0, "Sight L/C": 0.008, "D/P": 0.0015, "D/A": 0.0025}
    fee_key = next((k for k in pay_fees if k in payment), "사전 송금")
    total += base_price * pay_fees.get(fee_key, 0)
    fta_rates = {"협정 미적용 (기본세율)": 0.18, "한-미 FTA (KOR-USA)": 0.10, "한-EU FTA (KOR-EU)": 0.10, "한-중 FTA (KOR-CHINA)": 0.14, "RCEP": 0.12}
    if term == "DDP": total += base_price * fta_rates.get(fta_type, 0.18)
    return total

# --- [6. Plotly 스타일 차트 함수] ---
def draw_styled_chart(df, label):
    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=df['날짜'], y=df['환율'],
        mode='lines+markers',
        line=dict(color='#3d5afe', width=3, shape='spline'),
        marker=dict(size=6, color='white', line=dict(width=2, color='#3d5afe')),
        name=label
    ))
    fig.update_layout(
        title=dict(text=f"<b>{label} 추이 (최근 30일)</b>", font=dict(family='Pretendard', size=18, color='#1e293b')),
        template='plotly_white',
        margin=dict(l=20, r=20, t=60, b=20),
        height=350,
        hovermode='x unified',
        font=dict(family='Pretendard'),
        xaxis=dict(showline=True, linewidth=1, linecolor='lightgrey', gridcolor='#f1f5f9'),
        yaxis=dict(tickformat=',.2f', gridcolor='#f1f5f9', title="환율 (KRW)")
    )
    return fig

# --- [7. 서류 생성 함수] ---
def create_ci_docx(data):
    doc = Document(); doc.add_heading('COMMERCIAL INVOICE', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=6, cols=2); table.style = 'Table Grid'
    fields = [(f"① Shipper/Seller:\n{data['shipper']}", f"⑦ Invoice No. and date:\n{data['inv_no_date']}"),
              (f"② Consignee:\n{data['consignee']}", f"⑧ L/C No. and date:\n{data['lc_no_date']}"),
              (f"⑨ Buyer:\n{data['buyer']}", f"⑪ Terms: {data['terms']} / {data['transport']}"),
              (f"③ Departure date: {data['dep_date']}", f"⑫ Insurance: {data['insurance']}"),
              (f"④ Vessel: {data['vessel']} / From: {data['from_port']}", f"⑥ To: {data['to_port']}"),
              (f"⑬ FTA Agreement: {data['fta']}", f"⑭ Payment: {data['pay']}")]
    for i, (left, right) in enumerate(fields):
        table.rows[i].cells[0].text = left; table.rows[i].cells[1].text = right
    item_table = doc.add_table(rows=2, cols=6); item_table.style = 'Table Grid'
    for i, txt in enumerate(['Marks', 'Pkgs', 'Description', 'Qty', 'Price', 'Amount']): item_table.rows[0].cells[i].text = txt
    row = item_table.rows[1].cells
    row[0].text, row[1].text, row[2].text, row[3].text, row[4].text, row[5].text = data['marks'], data['pkg_kind'], data['description'], str(data['qty']), str(data['unit_price']), str(data['amount'])
    return doc

def create_pl_docx(data):
    doc = Document(); doc.add_heading('PACKING LIST', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=4, cols=2); table.style = 'Table Grid'
    table.rows[0].cells[0].text = f"Seller: {data['shipper']}"; table.rows[0].cells[1].text = f"Inv No: {data['inv_no_date']}"
    table.rows[1].cells[0].text = f"Consignee: {data['consignee']}"; table.rows[1].cells[1].text = f"Buyer: {data['buyer']}"
    item_table = doc.add_table(rows=2, cols=6); item_table.style = 'Table Grid'
    for i, txt in enumerate(['Marks', 'Pkgs', 'Goods', 'N.W', 'G.W', 'Meas']): item_table.rows[0].cells[i].text = txt
    row = item_table.rows[1].cells
    for i, key in enumerate(['marks', 'pkg_kind', 'description', 'net_weight', 'gross_weight', 'measure']): row[i].text = str(data[key])
    return doc

def create_bl_docx(data):
    doc = Document(); doc.add_heading('BILL OF LADING', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=3, cols=2); table.style = 'Table Grid'
    table.rows[0].cells[0].text = f"Shipper: {data['shipper']}"; table.rows[0].cells[1].text = f"B/L No: {data['bl_no']}"
    table.rows[1].cells[0].text = f"Consignee: {data['consignee']}"; table.rows[1].cells[1].text = f"Vessel: {data['vessel']}"
    table.rows[2].cells[0].text = f"Loading: {data['from_port']}"; table.rows[2].cells[1].text = f"Discharge: {data['to_port']}"
    return doc

# --- [8. 사이드바 구성] ---
with st.sidebar:
    st.title("💰 금융 & FTA 현황")
    current_rates = st.session_state['exchange_rates']
    st.metric(label=f"USD/KRW ({datetime.now().strftime('%Y-%m-%d')})", value=f"{current_rates['USD']:,.2f}원")
    st.markdown("---")
    st.subheader("🧮 환율 도구")
    with st.popover("🔍 간이 계산기 열기", use_container_width=True):
        st.markdown("### 🧮 실시간 환산")
        calc_currency = st.selectbox("전환 통화", ["USD", "JPY", "EUR", "CNY"], key="side_calc_curr")
        input_amt = st.number_input(f"{calc_currency} 금액", value=1000.0, key="side_calc_amt")
        rate = current_rates[calc_currency]
        krw_result = input_amt * (rate / 100) if calc_currency == "JPY" else input_amt * rate
        st.divider(); st.success(f"**결과:** {krw_result:,.0f} KRW")
    st.markdown("---")
    st.subheader("⚙️ 데이터 제어")
    if st.button("🔄 실시간 데이터 동기화"):
        with st.spinner("최신 환율 정보를 불러오는 중..."):
            new_rates = get_realtime_exchange_rates()
            st.session_state['exchange_rates'] = new_rates
            st.session_state['use_realtime'] = True
            st.success("API 동기화 완료")
            time.sleep(0.5); st.rerun()
    st.info("💡 동기화 시 yfinance API를 연동합니다.")

# --- [9. 메인 화면 로직] ---
st.title("🚢 Trade Master 2026: FTA & 결제 통합 자동화")

# 데이터 동적 로드
exchange_rates = st.session_state['exchange_rates']
use_realtime = st.session_state['use_realtime']
today_date = datetime.now().date()

# [브라우저용 디자인 테이블 생성 - 국기 아이콘 추가]
st.subheader(f"💵 주요 통화 환율 ({'실시간 API' if use_realtime else '2026 시뮬레이션'} 기반)")

# HTML 테이블 빌더
rates_html = f"""
<div class="modern-table-container">
    <table class="modern-table">
        <thead>
            <tr>
                <th>통화명</th>
                <th>매매기준율</th>
                <th>송금 보낼 때</th>
                <th>송금 받을 때</th>
            </tr>
        </thead>
        <tbody>
            <tr>
                <td class="currency-name">🇺🇸 USD 미국💲</td>
                <td class="rate-val">{exchange_rates['USD']:,.2f}</td>
                <td class="rate-diff diff-up">{exchange_rates['USD']*1.01:,.2f}</td>
                <td class="rate-diff diff-down">{exchange_rates['USD']*0.99:,.2f}</td>
            </tr>
            <tr>
                <td class="currency-name">🇯🇵 JPY 일본(100엔)💴</td>
                <td class="rate-val">{exchange_rates['JPY']:,.2f}</td>
                <td class="rate-diff diff-up">{exchange_rates['JPY']*1.01:,.2f}</td>
                <td class="rate-diff diff-down">{exchange_rates['JPY']*0.99:,.2f}</td>
            </tr>
            <tr>
                <td class="currency-name">🇪🇺 EUR 유럽💶</td>
                <td class="rate-val">{exchange_rates['EUR']:,.2f}</td>
                <td class="rate-diff diff-up">{exchange_rates['EUR']*1.01:,.2f}</td>
                <td class="rate-diff diff-down">{exchange_rates['EUR']*0.99:,.2f}</td>
            </tr>
            <tr>
                <td class="currency-name">🇨🇳 CNY 중국</td>
                <td class="rate-val">{exchange_rates['CNY']:,.2f}</td>
                <td class="rate-diff diff-up">{exchange_rates['CNY']*1.01:,.2f}</td>
                <td class="rate-diff diff-down">{exchange_rates['CNY']*0.99:,.2f}</td>
            </tr>
        </tbody>
    </table>
</div>
"""
st.markdown(rates_html, unsafe_allow_html=True)

# --- [Plotly 차트 섹션] ---
st.subheader("📈 주요 통화별 최근 30일 추이")
g_col1, g_col2 = st.columns(2)
currency_list = [("USD/KRW", "KRW=X", exchange_rates['USD'], 1), ("JPY/KRW (100엔)", "JPYKRW=X", exchange_rates['JPY'], 100),
                 ("EUR/KRW", "EURKRW=X", exchange_rates['EUR'], 1), ("CNY/KRW", "CNYKRW=X", exchange_rates['CNY'], 1)]

for i, (label, ticker, base, mult) in enumerate(currency_list):
    target_col = g_col1 if i % 2 == 0 else g_col2
    with target_col:
        df_hist = get_currency_history(ticker, base, mult, use_realtime, today_date)
        if not df_hist.empty: st.plotly_chart(draw_styled_chart(df_hist, label), use_container_width=True)

st.divider()
st.subheader("📑 거래 상세 및 가격 조건 설정")
with st.form("trade_form"):
    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown("**1. 기본 거래 정보**")
        shipper = st.text_area("수출자(Shipper)", "GILDING TRADING CO., LTD.\nSEOUL, KOREA")
        consignee = st.text_area("수입자(Consignee)", "MONARCH PRO CO., LTD.\nDETROIT, USA")
        from_port = st.text_input("출발지", "BUSAN, KOREA"); to_port = st.text_input("도착지", "DETROIT, USA")
        vessel = st.text_input("선박/항공편명", "PHEONIC V.123")
    with c2:
        st.markdown("**2. 인코텀즈 및 FTA**")
        selected_term = st.selectbox("Incoterms 2020", ["EXW", "FOB", "CIF", "DDP", "DAP", "CIP"])
        selected_fta = st.selectbox("FTA 협정 선택", ["협정 미적용 (기본세율)", "한-미 FTA (KOR-USA)", "한-EU FTA (KOR-EU)", "한-중 FTA (KOR-CHINA)", "RCEP"])
        transport_mode = st.radio("운송 수단", ["해상(SEA)", "항공(AIR)"], horizontal=True)
        insurance_type = st.selectbox("적하보험 조건", ["선택 안함", "ICC(A) (=ICC(AIR))", "ICC(B)", "ICC(C)"])
    with c3:
        st.markdown("**3. 품목 및 결제 정보**")
        payment = st.selectbox("결제방식", ["사전 송금 (Advance Payment)", "Sight L/C", "D/P", "D/A"])
        description = st.text_input("품명", "NYLON OXFORD")
        qty_input = st.number_input("수량", value=60000)
        selected_currency = st.selectbox("거래 통화", ["USD", "JPY", "EUR", "CNY"])
        unit_price_input = st.number_input(f"단가({selected_currency})", value=1.0 if selected_currency != "JPY" else 100.0)
    st.divider()
    subtotal = qty_input * unit_price_input
    estimated_total = calculate_estimated_cost(subtotal, selected_term, transport_mode, insurance_type, payment, selected_fta)
    final_rate = exchange_rates[selected_currency]
    total_krw = estimated_total * (final_rate / 100) if selected_currency == "JPY" else estimated_total * final_rate
    st.markdown(f"""<div class="info-box">💡 <b>최신 {selected_currency} 환율 반영 예상 총액:</b> {selected_currency} {estimated_total:,.2f} (약 {total_krw:,.0f} 원)</div>""", unsafe_allow_html=True)
    submitted = st.form_submit_button("🚀 분석 및 서류 생성")

if submitted:
    now = datetime(2026, 1, 30); formatted_inv_date = now.strftime('%b. %d. %Y').upper()
    data = {"shipper": shipper, "consignee": consignee, "from_port": from_port, "to_port": to_port, "vessel": vessel,
            "inv_no_date": f"INV-{now.year}-{now.strftime('%m%d')}\n{formatted_inv_date}", "lc_no_date": "LC-2026-001", "terms": selected_term, "transport": transport_mode,
            "insurance": insurance_type, "pay": payment, "fta": selected_fta, "description": description, "qty": f"{qty_input:,}", "unit_price": f"{unit_price_input:.2f}", 
            "amount": f"{estimated_total:,.2f} ({selected_currency})", "pkg_kind": "53 C/NO", "net_weight": "1,200 KGS", "gross_weight": "1,208 KGS", "marks": "MON/T DETROIT", "measure": "5.8 CBM", "bl_no": f"BK-{now.strftime('%y%m%d')}",
            "dep_date": (now + timedelta(days=7)).strftime('%b. %d. %Y').upper(), "buyer": consignee, "other_ref": "KOREA"}
    st.session_state['current_data'] = data
    with st.spinner("AI 관세사가 FTA 분석 중..."):
        risk_prompt = f"전문 관세사 분석: 통화 {selected_currency}, FTA {selected_fta}, 인코텀즈 {selected_term}, 결제 {payment}. PSR 충족 가능성과 대금 리스크를 한글로 분석하세요."
        response = client.chat.completions.create(model="gpt-4o", messages=[{"role": "user", "content": risk_prompt}])
        st.session_state['ai_analysis'] = response.choices[0].message.content

if 'ai_analysis' in st.session_state:
    t1, t2 = st.tabs(["💡 AI 전략 가이드", "📥 서류 다운로드"])
    with t1: st.markdown(st.session_state['ai_analysis'])
    with t2:
        curr = st.session_state['current_data']
        doc_files = {"Commercial_Invoice.docx": create_ci_docx(curr), "Packing_List.docx": create_pl_docx(curr), "Bill_of_Lading.docx": create_bl_docx(curr)}
        cols = st.columns(3)
        for i, (name, doc) in enumerate(doc_files.items()):
            bio = BytesIO(); doc.save(bio)
            cols[i].download_button(label=f"📥 {name}", data=bio.getvalue(), file_name=name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        st.success("모든 서류 생성이 완료되었습니다.")

